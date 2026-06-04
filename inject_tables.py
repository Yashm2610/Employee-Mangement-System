import docx
from docx.shared import Inches, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# Define the schema data
tables_data = {
    "1. employees (Core Employee Data)": [
        ["id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique internal database record identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL, UNIQUE", "Unique alphanumeric employee code."],
        ["emp_name", "VARCHAR(100)", "NOT NULL", "Full legal name of the employee."],
        ["email", "VARCHAR(100)", "NOT NULL", "Professional email address."],
        ["date_of_birth", "DATE", "NOT NULL", "Date of birth."],
        ["joining_date", "DATE", "NOT NULL", "Date of joining."],
        ["basic_salary", "DECIMAL(10,2)", "NOT NULL", "Monthly base salary payout."],
        ["age", "INT", "NOT NULL", "Employee's age."],
        ["gender", "VARCHAR(20)", "DEFAULT 'Male'", "Employee's gender."],
        ["education", "INT", "DEFAULT 2", "Educational qualification (0=High School, 1=Diploma, 2=Bachelor's, 3=Master's, 4=PhD)."],
        ["title", "VARCHAR(100)", "DEFAULT 'Software Engineer'", "Job designation/title."],
        ["department", "VARCHAR(100)", "DEFAULT 'General Affairs'", "Assigned department."],
        ["posting_location", "VARCHAR(100)", "DEFAULT 'Bangalore'", "City of employment."],
        ["payment_tier", "INT", "NOT NULL", "1=Executive, 2=Professional, 3=Associate."]
    ],
    "2. employee_bank_details (Bank Information)": [
        ["id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL, UNIQUE (FK)", "Foreign key mapping to employees."],
        ["bank_name", "VARCHAR(100)", "DEFAULT 'Bank of America'", "Name of the bank."],
        ["bank_account_num", "VARCHAR(50)", "DEFAULT '0000000000'", "Bank account number."],
        ["ifsc_code", "VARCHAR(20)", "DEFAULT 'BOFA0000001'", "IFSC Code."]
    ],
    "3. employee_financial_components (Unified Allowances/Deductions)": [
        ["id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL (FK)", "Foreign key mapping to employees."],
        ["component_name", "VARCHAR(100)", "NOT NULL", "Name of the allowance or deduction."],
        ["component_code", "TINYINT", "NOT NULL", "Type of component (1=Allowance, 2=Deduction)."],
        ["amount", "DECIMAL(12,2)", "DEFAULT 0.00", "Amount in INR."]
    ],
    "4. payslip_master (Payslip History)": [
        ["payslip_id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL (FK)", "Foreign key mapping to employees."],
        ["basic_salary", "DECIMAL(12,2)", "DEFAULT 0.00", "Base salary."],
        ["total_allowance", "DECIMAL(12,2)", "DEFAULT 0.00", "Total allowances."],
        ["total_deduction", "DECIMAL(12,2)", "DEFAULT 0.00", "Total deductions."],
        ["final_in_hand_salary", "DECIMAL(12,2)", "DEFAULT 0.00", "Final net salary paid to the employee."],
        ["generated_on", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Timestamp when generated."]
    ],
    "5. employee_holidays": [
        ["id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL (FK)", "Foreign key mapping to employees."],
        ["holiday_code", "TINYINT", "NOT NULL", "0=Present, 1=Casual Leave, 2=Sick Leave, 3=Paid Holiday, 4=Absent."]
    ],
    "6. employee_emails (Communication Log)": [
        ["id", "INT", "AUTO_INCREMENT PRIMARY KEY", "Unique identifier."],
        ["emp_id", "VARCHAR(50)", "NOT NULL (FK)", "Foreign key mapping to employees."],
        ["sender_email", "VARCHAR(100)", "DEFAULT 'admin@maxworth.com'", "Email address of the sender."],
        ["receiver_email", "VARCHAR(100)", "NULL", "Email address of the employee."],
        ["sent_at", "DATETIME", "DEFAULT CURRENT_TIMESTAMP", "Timestamp when email was sent."],
        ["subject", "VARCHAR(255)", "NULL", "Subject line of the email."],
        ["body", "TEXT", "NULL", "Body content of the email."],
        ["response_received_at", "DATETIME", "NULL", "Timestamp when employee responded."],
        ["status", "VARCHAR(20)", "DEFAULT 'Sent'", "Current status of the email interaction."]
    ]
}

def format_table(table):
    # Set header background to dark blue like original
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        shading_elm = parse_xml(r'<w:shd {} w:fill="2F3E46"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm)
        # Set text color to white
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True

def process_doc(filename):
    try:
        doc = docx.Document(filename)
        
        # We need to find "Database Schema Design"
        schema_idx = -1
        for i, p in enumerate(doc.paragraphs):
            if "Database Schema Design" in p.text:
                schema_idx = i
                break
                
        if schema_idx == -1:
            print(f"Heading not found in {filename}")
            return
            
        # Delete any paragraphs containing "-- =" or "CREATE TABLE" or "VARCHAR"
        to_delete = []
        for i in range(schema_idx + 1, len(doc.paragraphs)):
            text = doc.paragraphs[i].text
            if text.startswith("-- =") or "CREATE TABLE" in text or "VARCHAR" in text or "INT AUTO" in text or text.strip().endswith(");"):
                to_delete.append(doc.paragraphs[i])
                
        for p in to_delete:
            p._element.getparent().remove(p._element)

        # Now insert the new tables
        p_target = doc.paragraphs[schema_idx]
        
        # We'll build the tables in a new document, then copy the XML elements over
        # directly after p_target
        
        parent = p_target._element.getparent()
        index = parent.index(p_target._element) + 1
        
        # Add tables one by one
        for title, rows in tables_data.items():
            # Add title
            p_title = docx.text.paragraph.Paragraph(doc._element.body.add_p(), doc)
            p_title.add_run(title).bold = True
            p_title.style = 'Heading 3'
            parent.insert(index, p_title._element)
            index += 1
            
            # Create Table
            tbl = doc.add_table(rows=1, cols=4)
            tbl.style = 'Table Grid'
            hdr_cells = tbl.rows[0].cells
            hdr_cells[0].text = 'Field Name'
            hdr_cells[1].text = 'Data Type'
            hdr_cells[2].text = 'Constraints'
            hdr_cells[3].text = 'Description'
            
            # Style header
            for cell in hdr_cells:
                shading_elm = parse_xml(r'<w:shd {} w:fill="2F3E46"/>'.format(nsdecls('w')))
                cell._tc.get_or_add_tcPr().append(shading_elm)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.bold = True

            # Add rows
            for row_data in rows:
                row_cells = tbl.add_row().cells
                row_cells[0].text = row_data[0]
                row_cells[1].text = row_data[1]
                row_cells[2].text = row_data[2]
                row_cells[3].text = row_data[3]
                
            # Move the table element
            tbl_element = tbl._element
            tbl_element.getparent().remove(tbl_element)
            parent.insert(index, tbl_element)
            index += 1
            
            # Add an empty paragraph for spacing
            p_space = docx.text.paragraph.Paragraph(doc._element.body.add_p(), doc)
            parent.insert(index, p_space._element)
            index += 1

        doc.save(filename)
        print(f"Successfully injected tables into {filename}")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == '__main__':
    process_doc("Employee_Data_Management_System_Report.docx")
    process_doc("Employee_Data_Management_System_Report_Updated.docx")
    process_doc("Employee_Data_and_Payroll_System_Report.docx")

