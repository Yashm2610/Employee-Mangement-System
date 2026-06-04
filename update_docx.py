import docx

def update_schema(filename, sql_text):
    try:
        doc = docx.Document(filename)
        for i, p in enumerate(doc.paragraphs):
            if "Database Schema Design" in p.text:
                # Find the next table and delete it
                if doc.tables:
                    # Let's assume the first table after this paragraph is the schema table.
                    # Or we just delete all tables in the document? No, there might be other tables.
                    # python-docx doesn't easily link paragraphs and tables by index.
                    
                    # We can clear the table content instead or remove it using lxml
                    for tbl in doc.tables:
                        # Check if table has a cell with 'Field Name'
                        if tbl.rows and tbl.rows[0].cells and 'Field Name' in tbl.rows[0].cells[0].text:
                            # Delete the table
                            tbl._element.getparent().remove(tbl._element)
                            break
                
                # Insert the SQL text after the paragraph
                new_p = p.insert_paragraph_before("\n" + sql_text + "\n")
                new_p.style = 'Normal'
                break
                
        doc.save(filename)
        print(f"Successfully updated {filename}")
    except Exception as e:
        print(f"Error updating {filename}: {e}")

if __name__ == '__main__':
    with open('schema.sql', 'r') as f:
        sql_text = f.read()
    
    update_schema("Employee_Data_Management_System_Report.docx", sql_text)
    update_schema("Employee_Data_Management_System_Report_Updated.docx", sql_text)
    update_schema("Employee_Data_and_Payroll_System_Report.docx", sql_text)
